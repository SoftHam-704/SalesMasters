import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_paragraph('__________________________________________________')
        
        # Lists
        elif line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold text in list items
            parts = re.split(r'(\*\*.*?\*\*)', line[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Footer
        elif line.startswith('**SalesMasters 2026**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run('SalesMasters 2026')
            run.bold = True
        elif line.startswith('*Tecnologia a serviço'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line[1:-1])
            run.italic = True
            
        # Normal text
        else:
            # Handle bold text in paragraphs
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Styling
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.save(docx_path)
    print(f"File saved to: {docx_path}")

md_file = r"C:\Users\Systems\.gemini\antigravity\brain\89006149-1e14-47c0-b831-23d0604c2c1d\SalesMasters_Technical_Portfolio.md"
docx_file = r"C:\Users\Systems\.gemini\antigravity\brain\89006149-1e14-47c0-b831-23d0604c2c1d\SalesMasters_Technical_Portfolio.docx"

if __name__ == "__main__":
    convert_md_to_docx(md_file, docx_file)
